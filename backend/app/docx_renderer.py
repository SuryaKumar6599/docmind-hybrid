"""Stage 3: DOCX and PDF generation from TailoredContent.

Pure Python — NO LLM calls are made here. The LLM's job ended at Stage 2.
This module only structures and styles the text into a document.

PDF export uses pandoc when available, falling back to weasyprint, then
a plain-text .txt file with a warning.
"""
from __future__ import annotations

import logging
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .schemas import TailoredContent

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# DOCX builder
# ---------------------------------------------------------------------------

def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _section_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x31, 0x53, 0x43)  # DocMind moss green
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    # Underline separator
    p2 = doc.add_paragraph()
    run2 = p2.add_run("─" * 60)
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p2.paragraph_format.space_after = Pt(4)


def build_docx(
    content: TailoredContent,
    candidate_name: str,
    company: str,
    role: str,
    output_path: str | Path,
    original_resume_sections: dict[str, str] | None = None,
) -> Path:
    """Render *content* into a .docx file at *output_path*.

    Args:
        content: Stage 2 TailoredContent output.
        candidate_name: Full name for the document header.
        company: Target company name (used in cover letter).
        role: Target role (used in document title and cover letter).
        output_path: Where to write the .docx file.
        original_resume_sections: Optional dict with keys like
            'contact', 'education', 'skills' for pass-through sections.

    Returns:
        Path to the written .docx file.
    """
    doc = Document()

    # --- Title ---
    title = doc.add_heading(candidate_name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"Tailored for: {role} @ {company}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(10)
    subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # --- Professional Summary ---
    _section_heading(doc, "Professional Summary")
    doc.add_paragraph(content.tailored_summary)

    # --- Rewritten Experience ---
    _section_heading(doc, "Experience Highlights")
    sorted_bullets = sorted(content.rewritten_bullets, key=lambda b: b.priority)
    for bullet in sorted_bullets:
        _bullet(doc, bullet.rewritten)

    # --- Skills ---
    if content.skills_to_add:
        _section_heading(doc, "Additional Skills")
        doc.add_paragraph(", ".join(content.skills_to_add))

    # --- Pass-through sections (education, contact, etc.) ---
    if original_resume_sections:
        for section_title, section_body in original_resume_sections.items():
            if section_body.strip():
                _section_heading(doc, section_title)
                doc.add_paragraph(section_body.strip())

    # --- Cover Letter ---
    doc.add_page_break()
    _heading(doc, f"Cover Letter — {role} at {company}", level=1)
    doc.add_paragraph(content.cover_letter_opening)
    doc.add_paragraph(
        "I would welcome the opportunity to discuss how my experience aligns "
        f"with the {role} role at {company}. Thank you for your consideration."
    )
    closing = doc.add_paragraph("Sincerely,\n" + candidate_name)
    closing.paragraph_format.space_before = Pt(24)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    logger.info("DOCX written: %s", out)
    return out


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------

def _try_pandoc(docx_path: Path, pdf_path: Path) -> bool:
    """Convert docx → pdf via pandoc. Returns True on success."""
    if not shutil.which("pandoc"):
        return False
    try:
        subprocess.run(
            ["pandoc", str(docx_path), "-o", str(pdf_path), "--pdf-engine=xelatex"],
            check=True,
            capture_output=True,
            timeout=60,
        )
        logger.info("PDF written via pandoc: %s", pdf_path)
        return True
    except Exception as exc:
        logger.warning("pandoc failed: %s", exc)
        return False


def _try_weasyprint(docx_path: Path, pdf_path: Path) -> bool:
    """Convert docx → pdf via python-docx → HTML → weasyprint. Returns True on success."""
    try:
        import mammoth  # type: ignore[import-untyped]
        from weasyprint import HTML  # type: ignore[import-untyped]

        with open(docx_path, "rb") as f:
            result = mammoth.convert_to_html(f)
        html = result.value
        HTML(string=html).write_pdf(str(pdf_path))
        logger.info("PDF written via weasyprint: %s", pdf_path)
        return True
    except ImportError:
        logger.warning("weasyprint/mammoth not available — skipping")
        return False
    except Exception as exc:
        logger.warning("weasyprint failed: %s", exc)
        return False


def export_pdf(docx_path: Path | str) -> Path | None:
    """Export a .docx file to PDF. Returns the PDF path or None on failure."""
    docx_path = Path(docx_path)
    pdf_path = docx_path.with_suffix(".pdf")

    if _try_pandoc(docx_path, pdf_path):
        return pdf_path
    if _try_weasyprint(docx_path, pdf_path):
        return pdf_path

    logger.error(
        "PDF generation failed for %s. Install pandoc or weasyprint.", docx_path
    )
    return None


# ---------------------------------------------------------------------------
# Convenience: build both artifacts
# ---------------------------------------------------------------------------

def render_tailored_resume(
    content: TailoredContent,
    candidate_name: str,
    company: str,
    role: str,
    output_dir: str | Path,
    original_resume_sections: dict[str, str] | None = None,
) -> dict[str, str | None]:
    """Build DOCX and attempt PDF export. Returns {docx: path, pdf: path|None}."""
    out_dir = Path(output_dir)
    safe_name = f"{candidate_name.replace(' ', '_')}_{company}_{role}".replace("/", "-")
    docx_path = out_dir / f"{safe_name}.docx"

    docx = build_docx(content, candidate_name, company, role, docx_path, original_resume_sections)
    pdf = export_pdf(docx)
    return {"docx": str(docx), "pdf": str(pdf) if pdf else None}
